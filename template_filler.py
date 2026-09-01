# -*- coding: utf-8 -*-
"""
Word / Excel 模板填充。

约定：模板中使用 {{key}} 占位符，程序把对应值替换进去。
  - 文本字段：{{project_name}} 等，替换为对应文字。
  - 图片字段：{{铭牌图片}} 等，替换为上传的图片（插入到该位置）。
  - 附件字段：{{附件}}，替换为上传的附件（图片直接插入，其它文件记录文件名）。

Word 模板中的占位符可以放在正文段落或表格单元格里，两者都会被处理。
"""

import os
import re
from copy import copy


def _load_docx_api():
    """按需加载 Word 依赖，避免程序启动时提前加载 python-docx。"""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Cm
    return Document, qn, Cm


class TemplateFiller:
    """模板填充器（Word 用 python-docx，Excel 用 openpyxl）。"""

    # ---------- 文本占位符替换 ----------
    @staticmethod
    def _replace_in_paragraph(paragraph, placeholder, value):
        """替换单个段落里的占位符，兼容占位符跨多个 run 的情况。"""
        if placeholder not in paragraph.text:
            return False

        # 优先单 run 内替换，尽量保留原有格式
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
                return True

        # 占位符被拆分到多个 run 时：合并到第一个 run，清空其余
        full = "".join(r.text for r in paragraph.runs)
        new_full = full.replace(placeholder, value)
        if paragraph.runs:
            paragraph.runs[0].text = new_full
            for r in paragraph.runs[1:]:
                r.text = ""
        return True

    @staticmethod
    def _iter_paragraphs(doc):
        """遍历正文所有段落 + 所有表格单元格里的段落。"""
        for p in doc.paragraphs:
            yield p
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    # ---------- 图片插入 ----------
    @staticmethod
    def _insert_images_at(paragraph, images, max_width_cm=None, max_height_cm=None):
        """
        在指定段落末尾依次插入图片。

        images: [ (图片路径, 宽度cm), ... ]
        """
        _, _, cm = _load_docx_api()
        try:
            from PIL import Image
        except Exception:
            Image = None
        for index, (path, width_cm) in enumerate(images):
            if path and os.path.isfile(path):
                if index:
                    paragraph.add_run().add_break()
                try:
                    target_width = float(width_cm)
                except (TypeError, ValueError):
                    target_width = float(max_width_cm or 16.0)
                if max_width_cm:
                    target_width = min(target_width, float(max_width_cm))
                target_width = max(target_width, 1.0)
                picture_kwargs = {"width": cm(target_width)}
                if Image is not None:
                    try:
                        with Image.open(path) as image:
                            pixel_width, pixel_height = image.size
                        if pixel_width and pixel_height:
                            target_height = target_width * pixel_height / pixel_width
                            if max_height_cm and target_height > float(max_height_cm):
                                target_height = float(max_height_cm)
                                target_width = target_height * pixel_width / pixel_height
                                picture_kwargs["width"] = cm(target_width)
                            picture_kwargs["height"] = cm(target_height)
                    except Exception:
                        pass
                run = paragraph.add_run()
                run.add_picture(path, **picture_kwargs)

    @staticmethod
    def _replace_from_offset(paragraph, start, value):
        """从段落的字符位置 start 起替换到段落末尾，保留前面的文字样式。"""
        value = str(value or "")
        runs = list(paragraph.runs)
        if not runs:
            if value:
                paragraph.add_run(value)
            return

        # 起点恰好落在标签 run 末尾时，优先使用下一个 value run；这样标签可继续
        # 保持粗体，而冒号后的自动填充内容明确使用非粗体。
        offset = 0
        target_index = None
        target_offset = 0
        for index, run in enumerate(runs):
            text = run.text or ""
            end = offset + len(text)
            if start < end:
                target_index = index
                target_offset = max(0, start - offset)
                break
            if start == end:
                target_index = index + 1 if index + 1 < len(runs) else index
                target_offset = 0 if index + 1 < len(runs) else len(text)
                break
            offset = end

        if target_index is None:
            target_index = len(runs) - 1
            target_offset = len(runs[-1].text or "")

        for index, run in enumerate(runs):
            if index < target_index:
                continue
            if index == target_index:
                text = run.text or ""
                run.text = text[:target_offset] + value
                if value:
                    run.bold = False
            else:
                run.text = ""

    @staticmethod
    def _replace_paragraph_text(paragraph, value):
        """替换整段文字，保留第一段 run 的字体/段落版式。"""
        value = str(value or "")
        if paragraph.runs:
            paragraph.runs[0].text = value
            for run in paragraph.runs[1:]:
                run.text = ""
        elif value:
            paragraph.add_run(value)

    @staticmethod
    def _binding_spec(spec):
        """兼容旧版「标签: 字段名」和新版带位置模式的配置。"""
        if isinstance(spec, str):
            return {"key": spec, "mode": "same_line"}
        if isinstance(spec, dict) and spec.get("key"):
            return {"key": spec["key"], "mode": spec.get("mode", "same_line")}
        return None

    @staticmethod
    def _word_date(value):
        """把日期输入统一成 Word 中易读的中文日期格式。"""
        text = str(value or "").strip()
        match = re.search(r"(\d{4})\D+(\d{1,2})\D+(\d{1,2})", text)
        if match:
            return "%s年%s月%s日" % (
                match.group(1), int(match.group(2)), int(match.group(3)))
        return text

    @staticmethod
    def _replace_word_bindings(doc, bindings, text_map):
        """按模板标签及段落位置替换内容，不依赖高亮颜色。"""
        if not bindings:
            return

        specs = []
        for prefix, raw_spec in bindings.items():
            spec = TemplateFiller._binding_spec(raw_spec)
            if spec:
                specs.append((prefix, spec))
        specs.sort(key=lambda item: len(item[0]), reverse=True)
        paragraphs = list(TemplateFiller._iter_paragraphs(doc))

        def is_bound_label(text):
            clean = (text or "").strip()
            return any(prefix != "date" and clean.startswith(prefix)
                       for prefix, _ in specs)

        for index, paragraph in enumerate(paragraphs):
            text = paragraph.text or ""
            clean = text.strip()
            matched = None
            for prefix, spec in specs:
                if prefix != "date" and clean.startswith(prefix):
                    matched = (prefix, spec)
                    break
            if not matched:
                continue

            prefix, spec = matched
            value = text_map.get(spec["key"], "")
            mode = spec["mode"]
            if mode == "next_paragraph":
                # 工程动向标题本身独占一段，内容在它后面的下一段。
                for next_paragraph in paragraphs[index + 1:]:
                    next_text = next_paragraph.text or ""
                    if not next_text.strip():
                        continue
                    if is_bound_label(next_text):
                        break
                    TemplateFiller._replace_paragraph_text(next_paragraph, value)
                    break
            elif mode == "replace_paragraph":
                TemplateFiller._replace_paragraph_text(paragraph, value)
            else:
                # 同一行保留标签及冒号，只替换其后的内容。
                start = text.find(prefix) + len(prefix)
                while start < len(text) and text[start] in " \t:：":
                    start += 1
                TemplateFiller._replace_from_offset(paragraph, start, value)

        date_spec = next((spec for prefix, spec in specs if prefix == "date"), None)
        if date_spec:
            value = TemplateFiller._word_date(text_map.get(date_spec["key"], ""))
            date_pattern = re.compile(r"\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日")
            for paragraph in paragraphs:
                match = date_pattern.search(paragraph.text or "")
                if match:
                    TemplateFiller._replace_from_offset(
                        paragraph, match.start(), value)
                    break

    @staticmethod
    def _clear_highlights(doc):
        """清除成品 Word 中用于标注模板字段的底色/高亮。"""
        _, qn, _ = _load_docx_api()
        roots = [doc.element]
        for section in doc.sections:
            roots.extend([
                section.header._element, section.footer._element,
                section.first_page_header._element,
                section.first_page_footer._element,
                section.even_page_header._element,
                section.even_page_footer._element,
            ])
        seen = set()
        for root in roots:
            if id(root) in seen:
                continue
            seen.add(id(root))
            for tag in (qn("w:highlight"), qn("w:shd")):
                for node in list(root.iter(tag)):
                    parent = node.getparent()
                    if parent is not None:
                        parent.remove(node)

    # ---------- Word 填充 ----------
    @classmethod
    def fill_word(cls, template_path, output_path, text_map,
                  image_map=None, attachment_images=None, attachment_text="",
                  word_bindings=None, highlight_bindings=None):
        """
        填充 Word 模板。

        参数:
            template_path: 模板 .docx 路径
            output_path:   输出 .docx 路径
            text_map:      {"key": "value", ...} 替换 {{key}}
            image_map:     {"{{占位符}}": [(图片路径, 宽度cm), ...], ...}
            attachment_images: [(图片路径, 宽度cm), ...] 插入到 {{附件}} 占位符
            attachment_text: 非图片附件的文件名（写入 {{附件}} 位置）
        """
        Document, _, _ = _load_docx_api()
        doc = Document(template_path)

        # 图片按 Word 页面正文可用区域限制，避免照片超出页面宽度/高度。
        try:
            section = doc.sections[0]
            emu_per_cm = 360000.0
            max_width_cm = float(
                section.page_width - section.left_margin - section.right_margin
            ) / emu_per_cm
            max_height_cm = float(
                section.page_height - section.top_margin - section.bottom_margin
            ) / emu_per_cm
        except Exception:
            max_width_cm, max_height_cm = 16.0, 25.0
        if max_width_cm <= 0:
            max_width_cm = 16.0
        if max_height_cm <= 0:
            max_height_cm = 25.0

        # 1) 文本替换（正文 + 表格）
        for p in cls._iter_paragraphs(doc):
            for key, value in text_map.items():
                cls._replace_in_paragraph(p, "{{%s}}" % key, str(value if value is not None else ""))

        # 真实模板按标签和段落位置定位字段，不依赖模板底色。
        bindings = word_bindings if word_bindings is not None else highlight_bindings
        cls._replace_word_bindings(doc, bindings, text_map)

        # 2) 图片替换（铭牌图片等）
        for placeholder, images in (image_map or {}).items():
            for p in cls._iter_paragraphs(doc):
                if placeholder in p.text:
                    cls._replace_in_paragraph(p, placeholder, "")
                    cls._insert_images_at(p, images, max_width_cm, max_height_cm)
                    break  # 只处理第一处

        # 3) 附件插入（图片直接插入；非图片附件写入文件名）。
        # 即使用户没有上传附件，也要清掉模板占位符，避免它出现在成品里。
        for p in cls._iter_paragraphs(doc):
            if "{{附件}}" not in p.text:
                continue
            if attachment_images:
                cls._replace_in_paragraph(p, "{{附件}}", "")
                cls._insert_images_at(
                    p, attachment_images, max_width_cm, max_height_cm)
                if attachment_text:
                    p.add_run().add_break()
                    p.add_run(attachment_text)
            else:
                cls._replace_in_paragraph(p, "{{附件}}", attachment_text)
            break

        # 模板中的颜色只是编辑提示，成品中全部清除。
        cls._clear_highlights(doc)
        doc.save(output_path)
        return output_path

    # ---------- Excel 填充 ----------
    @classmethod
    def fill_excel(cls, template_path, output_path, text_map, sheet_name=None):
        """
        填充 Excel 模板（替换单元格里的 {{key}} 占位符）。

        参数:
            template_path: 模板 .xlsx 路径
            output_path:   输出 .xlsx 路径
            text_map:      {"key": "value", ...}
            sheet_name:    指定工作表名；为空则处理所有工作表
        """
        import openpyxl

        wb = openpyxl.load_workbook(template_path)
        sheets = [wb[sheet_name]] if sheet_name else wb.worksheets

        for ws in sheets:
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        for key, value in text_map.items():
                            placeholder = "{{%s}}" % key
                            if placeholder in cell.value:
                                cell.value = cell.value.replace(
                                    placeholder, str(value if value is not None else ""))

        wb.save(output_path)
        return output_path

    @classmethod
    def _is_total_sheet(cls, title, total_sheet_name):
        title = str(title or "")
        return bool(total_sheet_name and title == total_sheet_name) or "总表" in title

    @staticmethod
    def _safe_sheet_title(value):
        """把设备主人转换为 Excel 可接受的工作表名。"""
        title = re.sub(r'[\\/*?:\[\]]', "-", str(value or "").strip())
        title = re.sub(r"[\x00-\x1f]", "", title).strip(" .")
        return title[:31]

    @classmethod
    def _sheet_for_owner(cls, wb, owner_name, preferred_name, total_sheet_name):
        """按设备主人找到工作表；没有时复制一个非总表人员工作表的格式。"""
        total_sheet_name = str(total_sheet_name or "危险源统计表（总表）")
        total_titles = {
            ws.title for ws in wb.worksheets
            if cls._is_total_sheet(ws.title, total_sheet_name)
        }
        owner = str(owner_name or "").strip()

        if owner and owner in wb.sheetnames and owner not in total_titles:
            return wb[owner]

        if not owner:
            if preferred_name and preferred_name in wb.sheetnames:
                preferred = wb[preferred_name]
                if preferred.title not in total_titles:
                    return preferred
            for ws in wb.worksheets:
                if ws.title not in total_titles and ws.title != "Sheet1":
                    return ws
            return wb.active

        title = cls._safe_sheet_title(owner) or "未命名设备主人"
        if title in total_titles:
            title = "{}-分表".format(title)
        base_title = title
        suffix = 2
        while title in wb.sheetnames:
            tail = "-{}".format(suffix)
            title = base_title[:31 - len(tail)] + tail
            suffix += 1

        source = None
        if preferred_name and preferred_name in wb.sheetnames:
            candidate = wb[preferred_name]
            if candidate.title not in total_titles and candidate.title != "Sheet1":
                source = candidate
        if source is None:
            source = next(
                (ws for ws in wb.worksheets
                 if ws.title not in total_titles and ws.title != "Sheet1"),
                wb.active,
            )

        # copy_worksheet 会保留列宽、合并单元格、边框、字体、行高和表头。
        # 新分表只保留第 1-3 行表头，避免把模板人员的历史记录复制过去。
        new_ws = wb.copy_worksheet(source)
        new_ws.title = title
        for row in new_ws.iter_rows(min_row=4):
            for cell in row:
                cell.value = None
        return new_ws

    @classmethod
    def append_excel_row(cls, template_path, output_path, text_map,
                          sheet_name, column_map, owner_name=None,
                          total_sheet_name="危险源统计表（总表）"):
        """按设备主人选择/创建工作表，并按第二、三行表头追加一条记录。"""
        import openpyxl

        wb = openpyxl.load_workbook(template_path)
        ws = cls._sheet_for_owner(
            wb, owner_name, sheet_name, total_sheet_name)

        # 模板采用第 2、3 行双层表头；优先使用第三行，空白时回退到第二行。
        headers = {}
        for col in range(1, ws.max_column + 1):
            parts = []
            for row in (2, 3):
                value = ws.cell(row, col).value
                if value is not None and str(value).strip():
                    parts.append(str(value).strip())
            header = "\n".join(parts)
            if header:
                headers[header] = col
                headers.setdefault(str(ws.cell(3, col).value or "").strip(), col)

        # 在已有数据最后一行后写入，避免误覆盖历史记录。
        scan_columns = min(ws.max_column, 40)
        last_row = 3
        for row in range(4, ws.max_row + 1):
            if any(ws.cell(row, col).value not in (None, "")
                   for col in range(1, scan_columns + 1)):
                last_row = row
        target_row = last_row + 1

        # 复制上一条记录的样式、边框、批注和行高。
        style_row = last_row if last_row >= 4 else min(4, ws.max_row)
        ws.row_dimensions[target_row].height = ws.row_dimensions[style_row].height
        for col in range(1, ws.max_column + 1):
            src = ws.cell(style_row, col)
            dst = ws.cell(target_row, col)
            if src.has_style:
                dst._style = copy(src._style)
            if src.number_format:
                dst.number_format = src.number_format
            if src.alignment:
                dst.alignment = copy(src.alignment)
            if src.protection:
                dst.protection = copy(src.protection)

        for header, key in column_map.items():
            col = headers.get(header)
            if not col:
                continue
            value = text_map.get(key, "")
            # 让日期列保持可读文本，兼容 OCR 原始日期和手工输入。
            ws.cell(target_row, col).value = "" if value is None else str(value)

        # 序号按上一条记录递增；若历史编号为空则保留空白。
        serial_cell = ws.cell(target_row, 1)
        previous = 0
        if last_row >= 4:
            # 人员分表里有些历史记录序号为空，向上寻找最近一个数字再递增。
            for row in range(last_row, 3, -1):
                candidate = ws.cell(row, 1).value
                if candidate not in (None, ""):
                    previous = candidate
                    break
        try:
            serial_cell.value = int(previous) + 1
        except (TypeError, ValueError):
            serial_cell.value = ""

        wb.save(output_path)
        return output_path
