# -*- coding: utf-8 -*-
"""
本地自测脚本（在开发机上运行，不需要 OCR 模型）：
  1. 生成示例 Word / Excel 模板（含占位符）；
  2. 验证 Word 填充：文本替换 + 图片插入 + 附件插入；
  3. 验证 Excel 填充：单元格占位符替换；
  4. 通过 Flask 测试客户端完整走一遍「上传附件 → 生成文件」链路。

运行：python 测试/self_test.py
"""

import base64
import io
import os
import sys

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, BASE)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.utils import get_column_letter

from template_filler import TemplateFiller
from app import create_app

TPL_DIR = os.path.join(BASE, "templates")
OUT_DIR = os.path.join(BASE, "output")
os.makedirs(TPL_DIR, exist_ok=True)
os.makedirs(OUT_DIR, exist_ok=True)

# 1x1 像素红色 PNG（用于图片插入测试，无需 PIL）
TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)

TEST_MAP = {
    "serial_no": "LZD-2025-001",
    "date": "2025年9月1日",
    "project_name": "某某市某某道路改造工程",
    "owner": "某某市住房和城乡建设局",
    "builder": "某某建设集团有限公司",
    "supervisor": "某某监理咨询有限公司",
    "designer": "某某建筑设计研究院",
    "category": "框架结构",
    "area": "10510平方米",
    "start_date": "2025年9月1日",
    "end_date": "2026年8月31日",
    "manager": "张三",
    "manager_phone": "13800000001",
    "site_manager": "王五",
    "site_phone": "13800000002",
    "phone": "13800000000",
    "address": "某某市某某区某某路1号",
    "department": "工程部",
    "handler": "李四",
    "reason": "申请办理施工许可相关手续",
}


def make_word_template():
    """生成示例 Word 模板（正文 + 表格占位符 + 图片/附件占位符）。"""
    path = os.path.join(TPL_DIR, "流转单模板.docx")
    doc = Document()
    title = doc.add_heading("工程流转单", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("编号：{{serial_no}}")

    table = doc.add_table(rows=10, cols=4)
    table.style = "Table Grid"
    rows_data = [
        ["工程名称", "{{project_name}}", "建设地址", "{{address}}"],
        ["建设单位", "{{owner}}", "施工单位", "{{builder}}"],
        ["监理单位", "{{supervisor}}", "设计单位", "{{designer}}"],
        ["工程类别", "{{category}}", "建筑面积", "{{area}}"],
        ["开工日期", "{{start_date}}", "竣工日期", "{{end_date}}"],
        ["项目经理", "{{manager}}", "项目经理手机", "{{manager_phone}}"],
        ["文明施工专管员", "{{site_manager}}", "专管员手机", "{{site_phone}}"],
        ["联系电话", "{{phone}}", "流转单编号", "{{serial_no}}"],
        ["经办人", "{{handler}}", "申请部门", "{{department}}"],
        ["事由", "{{reason}}", "", ""],
    ]
    for r, row_data in enumerate(rows_data):
        for c, text in enumerate(row_data):
            table.cell(r, c).text = text

    doc.add_paragraph("")
    doc.add_paragraph("工程铭牌照片：")
    doc.add_paragraph("{{铭牌图片}}")
    doc.add_paragraph("")
    doc.add_paragraph("附件：")
    doc.add_paragraph("{{附件}}")
    doc.save(path)
    return path


def make_excel_template():
    """生成示例 Excel 台账模板。"""
    path = os.path.join(TPL_DIR, "台账模板.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "台账"
    headers = ["编号", "工程名称", "建设地址", "建设单位", "施工单位",
               "监理单位", "设计单位", "工程类别", "建筑面积", "开工日期",
               "竣工日期", "项目经理", "项目经理手机", "联系电话", "事由"]
    placeholders = ["{{serial_no}}", "{{project_name}}", "{{address}}", "{{owner}}",
                    "{{builder}}", "{{supervisor}}", "{{designer}}", "{{category}}",
                    "{{area}}", "{{start_date}}", "{{end_date}}", "{{manager}}",
                    "{{manager_phone}}", "{{phone}}", "{{reason}}"]
    for c, h in enumerate(headers, 1):
        ws.cell(row=1, column=c, value=h)
    for c, p in enumerate(placeholders, 1):
        ws.cell(row=2, column=c, value=p)
    for c in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(c)].width = 24
    wb.save(path)
    return path


def test_fill_word():
    tpl = os.path.join(TPL_DIR, "流转单模板.docx")
    out = os.path.join(OUT_DIR, "测试_流转单.docx")
    img = os.path.join(OUT_DIR, "_tiny.png")
    with open(img, "wb") as f:
        f.write(TINY_PNG)

    TemplateFiller.fill_word(
        tpl, out, TEST_MAP,
        image_map={"{{铭牌图片}}": [(img, 8.0)]},
        attachment_images=[(img, 8.0)],
        attachment_text="",
    )

    d = Document(out)
    texts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                texts.append(cell.text)
    all_text = "\n".join(texts)

    for key in TEST_MAP:
        assert "{{%s}}" % key not in all_text, "占位符未替换: %s" % key
    assert "{{铭牌图片}}" not in all_text, "图片占位符未清除"
    assert "{{附件}}" not in all_text, "附件占位符未清除"
    assert TEST_MAP["project_name"] in all_text
    assert len(d.inline_shapes) == 2, "应插入 2 张图片，实际 %d" % len(d.inline_shapes)
    print("  ✓ Word 填充：文本替换、图片/附件插入均通过")


def test_fill_excel():
    tpl = os.path.join(TPL_DIR, "台账模板.xlsx")
    out = os.path.join(OUT_DIR, "测试_台账.xlsx")
    TemplateFiller.fill_excel(tpl, out, TEST_MAP)
    wb = openpyxl.load_workbook(out)
    ws = wb.active
    vals = []
    for row in ws.iter_rows():
        for cell in row:
            vals.append(cell.value)
    joined = " ".join(str(v) for v in vals if v is not None)
    for key in TEST_MAP:
        assert "{{%s}}" % key not in joined, "Excel 占位符未替换: %s" % key
    assert TEST_MAP["project_name"] in joined
    print("  ✓ Excel 填充：单元格占位符替换通过")


def test_api_generate():
    app = create_app()
    app.config["TESTING"] = True
    client = app.test_client()

    # 模拟上传附件（图片）
    resp = client.post("/api/upload_attachment",
                       data={"file": (io.BytesIO(TINY_PNG), "盖章扫描件.png")},
                       content_type="multipart/form-data")
    assert resp.status_code == 200
    att = resp.get_json()
    assert att["ok"] is True

    # 模拟生成（未传铭牌图片，验证字段+附件链路）
    payload = {
        "fields": TEST_MAP,
        "image_name": None,
        "attachments": [{"name": att["filename"], "original": att["original"]}],
    }
    resp = client.post("/api/generate", json=payload)
    assert resp.status_code == 200, resp.get_data(as_text=True)
    data = resp.get_json()
    assert data["ok"] is True
    assert len(data["outputs"]) == 2, data
    for o in data["outputs"]:
        assert os.path.isfile(os.path.join(OUT_DIR, o["name"]))
    print("  ✓ 生成接口：Word + Excel 完整链路通过")
    print("    生成文件：" + "、".join(o["name"] for o in data["outputs"]))


def test_config_endpoint():
    app = create_app()
    client = app.test_client()
    resp = client.get("/api/config")
    assert resp.status_code == 200
    cfg = resp.get_json()
    assert cfg["fields"], "config 未返回字段"
    print("  ✓ 配置接口：字段/模板配置读取通过")


if __name__ == "__main__":
    print("=" * 52)
    print("自测开始（生成示例模板 + 验证填充与生成链路）")
    make_word_template()
    make_excel_template()
    print("  已生成示例模板：templates/流转单模板.docx、templates/台账模板.xlsx")
    test_fill_word()
    test_fill_excel()
    test_api_generate()
    test_config_endpoint()
    print("=" * 52)
    print("全部测试通过 ✓")
